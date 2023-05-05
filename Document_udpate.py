#!/usr/bin/env python
# _*_ coding:utf-8 _*_
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.shape import CT_Inline
from WeekTime_parser import week_number, today, last_week, week_check_day
from Weekly_parser import ser_previous_scrap_table
from chart import ChartIOGen
from itertools import cycle
import yaml
import pandas as pd
import Template_generator as Tg


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def add_picture_to_run(run, picture_file, width=None, height=None):
    """
    Add a picture at the end of a run.
    """
    rid, image = run.part.get_or_add_image(picture_file)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = run.part.next_id, image.filename
    inline = CT_Inline.new_pic_inline(shape_id, rid, filename, cx, cy)
    run._r.add_drawing(inline)


def custom_style_set(doc, fontsize, bold=True, base='Normal', color='', alignment=''):
    try:
        if color == 'g':
            rgb_color = RGBColor(0, 176, 80)
        elif color == 'b':
            rgb_color = RGBColor(0, 0, 255)
        else:
            rgb_color = RGBColor(0, 0, 0)
            color = ''
        styles = doc.styles
        custom_style = styles.add_style('{}{}{}{}'.format('bold' if bold else 'normal', fontsize, color, alignment),
                                        WD_STYLE_TYPE.PARAGRAPH)
        custom_paragraph_format = custom_style.paragraph_format
        custom_paragraph_format.line_spacing = 1
        custom_style.base_style = styles[base]
        if alignment:
            if alignment == 'center':
                custom_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif alignment == 'right':
                custom_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif alignment == 'justify':
                custom_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        font = custom_style.font
        font.name = 'Tahoma'
        font.size = Pt(fontsize)
        font.bold = bold
        font.color.rgb = rgb_color
    except ValueError:  # style already exist
        pass


def custom_style_setup(doc):
    custom_style_set(doc, 10)  # bold10
    custom_style_set(doc, 12)  # bold12
    custom_style_set(doc, 14)  # bold14
    custom_style_set(doc, 12, bold=False, alignment='center')  # normal12center
    custom_style_set(doc, 12, bold=False, alignment='justify')  # normal12justify
    custom_style_set(doc, 20, bold=False, alignment='center')  # normal20center
    custom_style_set(doc, 12, color='b')  # bold12b
    custom_style_set(doc, 12, color='g')  # bold12g
    custom_style_set(doc, 12, alignment='center')  # bold12center
    custom_style_set(doc, 12, alignment='right')  # bold12right
    custom_style_set(doc, 12, alignment='justify')  # bold12justify
    custom_style_set(doc, 12, color='b', alignment='right')  # bold12bright
    custom_style_set(doc, 12, color='g', alignment='right')  # bold12gright


def df_to_table(doc, df):
    """
    support multi-index
    :param doc:
    :param df:
    :return:
    """
    # add a table to the end and create a reference variable
    # extra column is generate so we can add the index column
    index_cols = df.index.nlevels
    t = doc.add_table(df.shape[0], df.shape[1] + index_cols, style='summary table')

    # add the index columns.
    index_style_tup = ('normal20center', 'bold12center', 'normal12center')  # first column, second column, third column
    column_style_dict = {0: 'bold12gright', 1: 'bold12gright'}  # , 2: 'bold12bright'
    for i in range(index_cols):
        previous_value = None
        previous_loc = (0, i)
        style = index_style_tup[i]
        for j in range(df.shape[0]):
            this_value = df.index[j][i]
            # t.cell(j, i).text = this_value
            if this_value != previous_value:
                if previous_value is not None:
                    t.cell(*previous_loc).merge(t.cell(j - 1, i))
                previous_loc = (j, i)
                t.cell(j, i).text = this_value
                t.cell(j, i).paragraphs[
                    0].style = 'bold12center' if previous_value is None else style  # None mean first row
                t.cell(j, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            previous_value = this_value

        else:  # last row
            t.cell(*previous_loc).merge(t.cell(j, i))
            t.cell(j, i).text = this_value
            t.cell(j, i).paragraphs[0].style = style
            t.cell(j, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            # add the rest of the dataframe
    table_width = range(df.shape[1])
    table_length = range(df.shape[0])
    boolean_toggle = cycle([True, False])
    for i in table_length:
        this_value_is_pcs = next(boolean_toggle)
        for j in table_width:
            this_value = df.values[i, j]
            if not isinstance(this_value, str):
                # if isinstance(this_value, float):  # float64 is inherited from float
                #     this_value = '{:,.1%}'.format(this_value)
                # else:
                #     this_value = str(this_value)
                if this_value_is_pcs:
                    this_value = "{0:.0f}".format(this_value)
                else:
                    this_value = '{:,.1%}'.format(this_value)

            t.cell(i, j + index_cols).text = this_value
            if j != table_width[-1]:
                t.cell(i, j + index_cols).paragraphs[0].style = column_style_dict.get(j, 'bold12right')
            else:
                t.cell(i, j + index_cols).paragraphs[0].style = 'bold12bright'


def add_scrap_table(doc, date):
    # date_wt = copy.copy(date)
    table = doc.add_table(rows=1, cols=5, style='Table Grid')
    ser_docx_scarp_pcs = ser_previous_scrap_table()
    records = ('Scraped pcs', '0', '0', '0', '0')
    t, qty1, qty2, qty3, qty4 = records
    row_cells = table.add_row().cells
    row_cells[0].text = t
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Week'
    # AUTO scrap pcs
    row_num = 1
    for week, pcs in ser_docx_scarp_pcs.iteritems():
        row_cells[row_num].text = str(pcs)
        hdr_cells[row_num].text = week
        row_num += 1
    # AUTO scrap pcs
    # all 0 scrap pcs
    # row_cells[1].text = qty1
    # row_cells[2].text = qty2
    # row_cells[3].text = qty3
    # row_cells[4].text = qty4
    # for i in range(4):
    #     n = 4 - i
    #     hdr_cells[n].text = week_number(date_wt)
    #     date_wt = last_week(date_wt)
    # all 0 scrap pcs
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(table.columns):
        if i == len(table.columns) - 1:
            style = 'bold12b'
        else:
            style = 'bold12'
        for cell in col.cells:
            cell.width = Cm(3.2) if i == 0 else Cm(2.22)
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.style = doc.styles[style]


sbl_table_header = (
    'Defect type', 'PIC', 'Suffered map', 'Weekly suffered rate', "Q'ty", 'Suffered wafer', 'Suffered rate')


def add_sbl_table(doc, side):
    col_len = len(sbl_table_header)
    table = doc.add_table(rows=2, cols=col_len, style='{}_SBL'.format(side))
    hdr_cells = table.rows[0].cells
    for i in range(col_len):
        hdr_cells[i].text = sbl_table_header[i]
    for col in table.columns:
        for cell in col.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.style = doc.styles['bold10']


sides = ['FSV', 'BSV', 'CP']


def add_charts_data_for_part(doc, part):
    # part_have_bsl = part in Tg.df_ref_dict
    doc.add_paragraph('Weekly yield trend of {}'.format(part), style='square')
    week_yield = Tg.summary_cache.get((part, 'Yield'))
    if week_yield is None:
        doc.add_paragraph('No fab out', style='purple square')
    else:
        doc.add_paragraph('Final yield:{:,.1%}'.format(week_yield), style='purple square')
        weekly_yield_trend_chart = Tg.summary_cache[part, 'chart', 'Weekly Yield Trend']
        doc.add_picture(weekly_yield_trend_chart.ref,
                        width=Cm(19.3))  # image had been modified by openpyxl to openpyxl Image object
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # align picture to center

        doc.add_paragraph('Final yield breakdown:', style='diamond')
        for side in sides:
            loss = Tg.summary_cache[part, side]
            doc.add_paragraph('{} loss:{:,.2%}'.format(side, loss), style='check mark')
        weekly_final_yield_loss_breakdown_chart = Tg.summary_cache[part, 'chart', 'Weekly Final Yield Loss Breakdown']
        doc.add_picture(weekly_final_yield_loss_breakdown_chart.ref, width=Cm(19.3))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for side in sides:
            bsl = Tg.summary_cache.get((part, 'baseline', side))
            if bsl is None:
                bsl_str = ''
            else:
                bsl_str = ' (Baseline: {:,.2%})'.format(bsl)
            doc.add_paragraph('{} defect trend – {}{}'.format(side, part, bsl_str), style='diamond')
            # side_defect_highlight_dict = Tg.summary_cache[part, 'greater_than_bsl_highlight', side]
            side_defect_highlight_dict = Tg.summary_cache.get((part, 'greater_than_bsl_highlight', side), Tg.NestedDict())
            if side_defect_highlight_dict.dict:
                for defect, increase in side_defect_highlight_dict.dict.items():
                    doc.add_paragraph('{} YLD loss of {} is increased by {:,.2%}.'.format(side, defect, increase),
                                      style='check mark')
            else:
                doc.add_paragraph('{} YLD loss is stable.'.format(side), style='check mark')
            side_yield_loss_trend_chart = Tg.summary_cache[part, 'chart', '{} Yield Loss Trend'.format(side)]
            doc.add_picture(side_yield_loss_trend_chart.ref, width=Cm(19.3))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('{} YLD & YLD loss chart'.format(part), style='diamond')

        fsv_pareto = Tg.summary_cache[part, 'chart', 'FSV Yield Loss Pareto']
        bsv_pareto = Tg.summary_cache[part, 'chart', 'BSV Yield Loss Pareto']
        # doc.add_picture(fsv_pareto.ref, width=Cm(9.4)).add_run().add_picture(bsv_pareto.ref, width=Cm(9.4))
        run = doc.add_paragraph().add_run()
        add_picture_to_run(run, fsv_pareto.ref, width=Cm(9.4))
        add_picture_to_run(run, bsv_pareto.ref, width=Cm(9.4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for side in ['FSV', 'BSV']:
            doc.add_paragraph('{} {} > SBL OOC case  '.format(part, side), style='check mark underline')
            add_sbl_table(doc, side)


ooc_table_header = ('Part', 'Group', "Q'ty", 'Action')


def add_ooc_chart_table(doc):
    chart = ChartIOGen()
    ooc_chart, df_ooc_info = chart.weekly_ooc()
    ooc_str = str()
    for part in Tg.ooc_parts_show_in_doc:
        part = part + '/'
        ooc_str += part
    ooc_str.rstrip('/')
    doc.add_paragraph('{} OOC Case'.format(ooc_str), style='square')
    doc.add_picture(ooc_chart.ref)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    doc.add_paragraph('Conti-Gen1 OOC case', style='check mark')
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4, style='ooc_table')
    col_len = len(ooc_table_header)
    hdr_cells = table.rows[0].cells
    for i in range(col_len):
        hdr_cells[i].text = ooc_table_header[i]
    ooc_normal_count = 0
    for part in df_ooc_info.index:
        parse_str = df_ooc_info.at[part, 'groups']
        if parse_str == 'Normal':
            # table.add_row()
            ooc_normal_count += 1
        else:
            grp_dict = yaml.load('{' + parse_str + '}', Loader=yaml.BaseLoader)
            for group, count in grp_dict.items():
                row = table.add_row()
                cells = row.cells
                cells[0].text = part
                cells[1].text = group
                cells[2].text = str(count)
                cells[3].text = 'Suffer {}, will keep trace it.'.format(group)
    if ooc_normal_count == len(df_ooc_info.index):  # this week doesn't have ooc cases, add an empty row
        table.add_row()
    table_cell_width = (2.1, 6.1, 1.2, 9.2)
    for i1, col in enumerate(table.columns):
        for i2, cell in enumerate(col.cells):
            if i2 == 0:
                style = 'bold12justify'
            else:
                if i1 == 2:
                    style = 'normal12center'
                else:
                    style = 'normal12justify'
            cell.width = Cm(table_cell_width[i1])
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.style = doc.styles[style]


def report_gen(date, df_summary_table, filename=None):
    title_ele = 'Conti & SMP Pressure Sensor Weekly Report {} {}'
    if filename is None:
        path = Tg.dump_data_root_folder + '\\Weekly_Report\\{}_data'.format(week_check_day(date).year)
        file = '\\D632_Weekly_{}.docx'.format(week_number(week_check_day(date)))
        filename = path + file
    doc = Document(r'templates/template.docx')
    for paragraph in doc.paragraphs:
        delete_paragraph(paragraph)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.95)
        section.right_margin = Cm(0.75)
    custom_style_setup(doc)
    # title set
    t = doc.add_heading()
    t_paragraph_format = t.paragraph_format
    t_paragraph_format.line_spacing = 1
    t.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title = t.add_run(title_ele.format(week_number(week_check_day(date)), week_check_day(date).year))
    title.font.name = 'Tahoma'
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)

    # start writing content
    doc.add_paragraph('Key indices:', style='bold14')
    doc.add_paragraph()
    doc.add_paragraph('Scrap:', style='arrow bullet')
    doc.add_paragraph('Target: <0.5%', style='content')
    doc.add_paragraph()
    add_scrap_table(doc, date)
    doc.add_paragraph()
    doc.add_paragraph('CP/ overall yield:', style='arrow bullet')
    # add summary table
    df_to_table(doc, df_summary_table.loc[Tg.docx_summary_table_header])
    doc.add_page_break()
    # second page

    for part in Tg.docx_summary_charts_parts:
        add_charts_data_for_part(doc, part)
    add_ooc_chart_table(doc)
    doc.save(filename)


def main():
    df_ = pd.read_excel('W03.xlsx', sheet_name='Weekly_summary_table', index_col=0, header=[0, 1, 2])
    report_gen(today, df_.T, 'demo1.docx')


if __name__ == '__main__':
    main()
